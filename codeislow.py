#!/usr/bin/env python3


import datetime
from importlib import reload
import json
import os
import re
import sys
import time
from pathlib import Path

import docx
import requests
from PyPDF2 import PdfReader
from bottle import Bottle
from bottle import request, static_file, template, HTTPError, run
from bottle_sslify import SSLify
from dotenv import load_dotenv, find_dotenv
from odf import text, teletype
from odf.opendocument import load
from requests.adapters import HTTPAdapter
from urllib3.util import Retry
    

ARTICLE_REGEX = re.compile(r"(^.*(?:article|art\.).*$)", flags=re.I | re.M)
MAIN_CODELIST = {
    "CCIV": "Code civil",
    "CPRCIV": "Code de procédure civile",
    "CCOM": "Code de commerce",
    "CTRAV": "Code du travail",
    "CPI": "Code de la propriété intellectuelle",
    "CPEN": "Code pénal",
    "CPP": "Code de procédure pénale",
    "CASSUR": "Code des assurances",
    "CCONSO": "Code de la consommation",
    "CSI": "Code de la sécurité intérieure",
    "CSP": "Code de la santé publique",
    "CSS": "Code de la sécurité sociale",
    "CESEDA": "Code de l'entrée et du séjour des étrangers et du droit d'asile",
    "CGCT": "Code général des collectivités territoriales",
    "CPCE": "Code des postes et des communications électroniques",
    "CENV": "Code de l'environnement",
    "CJA": "Code de justice administrative",
}

REG_BEGGINNING = {
    "UNIVERSAL": r"((?:L\.?|R\.?|A\.?|D\.?)?\s*\d+-?\d*-?\d*),?\s*(?:alinéa|al\.)?\s*\d*\s?°?\s*"
                 r"(?:,\s*((?:L\.?|R\.?|A\.?|D\.?)?\s*\d+-?\d*-?\d*),?\s*(?:alinéa|al\.)?\s*\d*\s?°?\s*"
                 r"(?:,\s*((?:L\.?|R\.?|A\.?|D\.?)?\s*\d+-?\d*-?\d*),?\s*(?:alinéa|al\.)?\s*\d*\s?°?\s*"
                 r"(?:,\s*((?:L\.?|R\.?|A\.?|D\.?)?\s*\d+-?\d*-?\d*),?\s*(?:alinéa|al\.)?\s*\d*\s?°?\s*)?)?)?"
                 r"(?:et\s*((?:L\.?|R\.?|A\.?|D\.?)?\s*\d+-?\d*-?\d*),?\s*(?:alinéa|al\.)?\s*\d*\s?°?\s*)?"
                 r"(?:et s\.|et suivants)?",
}

REG_ENDING = {
    "CCIV": r"\s*(?:du Code civil|C\. civ\.)",
    "CPRCIV": r"\s*(?:du Code de procédure civile|C\. pr\. civ\.|CPC|du CPC)",
    "CCOM": r"\s*(?:du Code de commerce|C\. com\.)",
    "CTRAV": r"\s*(?:du Code du travail|C\. trav\.)",
    "CPI": r"\s*(?:du Code de la propriété intellectuelle|CPI|C\. pr\. int\.|du CPI)",
    "CPEN": r"\s*(?:du Code pénal|C\. pén\.)",
    "CPP": r"\s*(?:du Code de procédure pénale|du CPP|CPP)",
    "CASSUR": r"\s*(?:du Code des assurances|C\. assur\.)",
    "CCONSO": r"\s*(?:du Code de la consommation|C\. conso\.)",
    "CSI": r"\s*(?:du Code de la sécurité intérieure|CSI|du CSI)",
    "CSP": r"\s*(?:du Code de la santé publique|C\. sant\. pub\.|CSP|du CSP)",
    "CSS": r"\s*(?:du Code de la sécurité sociale|C\. sec\. soc\.|CSS|du CSS)",
    "CESEDA": r"\s*(?:du Code de l'entrée et du séjour des étrangers et du droit d'asile|CESEDA|du CESEDA)",
    "CGCT": r"\s*(?:du Code général des collectivités territoriales|CGCT|du CGCT)",
    "CPCE": r"\s*(?:du Code des postes et des communications électroniques|CPCE|du CPCE)",
    "CENV": r"\s*(?:du Code de l'environnement|C. envir.|CE |du CE )",
    "CJA": r"\s*(?:du Code de justice administrative|CJA|du CJA)",
}

CODES_REGEX = {
    "CCIV": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CCIV"],
    "CPRCIV": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CPRCIV"],
    "CCOM": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CCOM"],
    "CTRAV": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CTRAV"],
    "CPI": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CPI"],
    "CPEN": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CPEN"],
    "CPP": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CPP"],
    "CASSUR": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CASSUR"],
    "CCONSO": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CCONSO"],
    "CSI": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CSI"],
    "CSP": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CSP"],
    "CSS": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CSS"],
    "CESEDA": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CESEDA"],
    "CGCT": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CGCT"],
    "CPCE": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CPCE"],
    "CENV": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CENV"],
    "CJA": REG_BEGGINNING["UNIVERSAL"] + REG_ENDING["CJA"],

}

app = Bottle()


def legifrance_auth():
    '''Get auth token from SECRETS into dotenv'''
    token_url = "https://oauth.piste.gouv.fr/api/oauth/token"

    load_dotenv()
    client_id = os.environ.get("API_KEY")
    client_secret = os.environ.get("API_SECRET")
    
    if client_id is None or client_secret is None:
        # return HTTPError(401, "No credential have been set")
        raise Exception("No credential have been set")
    res = requests.post(
        token_url,
        data={
            "grant_type": "client_credentials",
            "client_id": client_id,
            "client_secret": client_secret,
            "scope": "openid",
        },
    )
    if res.status_code > 399:
        # return HTTPError(res.status_code, "Unauthorized: invalid credentials")
        raise Exception(f"HTTP Error code: {res.status_code}: Invalid credentials")
    response = res.json()
    token = response["access_token"]
    legifrance_headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
    }
    return legifrance_headers


def spaces_remover(string):
    return re.sub(" {2,}", " ", string)


# Ouverture du fichier utilisateur
# Indicateur de progression quand le fichier est suffisamment long
def file_opener(ext, file_path):
    
    paragraphsdoc = []
    previous_progress = int
    f = open(file_path, "rb")
    if ext == ".docx":
        document = docx.Document(f)
        for i in range(len(document.paragraphs)):
            if len(document.paragraphs) > 20:
                progress = round((i / len(document.paragraphs)) * 100)
                if progress % 10 == 0 and previous_progress != progress:
                    previous_progress = progress
                    yield str(progress) + " % ... "
            paragraph = document.paragraphs[i].text
            if ARTICLE_REGEX.search(paragraph) is not None:
                paragraphsdoc.append(paragraph)
    elif ext == ".odt":
        document = load(f)
        texts = document.getElementsByType(text.P)
        for i in range(len(texts)):
            if (len(texts)) > 20:
                progress = round((i / len(texts)) * 100)
                if progress % 10 == 0 and previous_progress != progress:
                    previous_progress = progress
                    yield str(progress) + " % ... "
            paragraph = teletype.extractText(texts[i])
            if ARTICLE_REGEX.search(paragraph) is not None:
                paragraphsdoc.append(paragraph)
    elif ext == ".pdf":
        reader = PdfReader(file_path)
        for i in range(len(reader.pages)):
            if (len(reader.pages)) > 10:
                progress = round((i / len(reader.pages)) * 100)
                if previous_progress != progress:
                    previous_progress = progress
                    yield str(progress) + " % ... "
            page = reader.pages[i]
            page_text = (page.extract_text())
            if ARTICLE_REGEX.search(page_text) is not None:
                paragraphsdoc.append(page_text)
    complete_text = spaces_remover(" ".join(paragraphsdoc))
    f.close()
    return complete_text


# Les paragraphes à tester sont confrontés à l'expression régulière de chaque code
def code_detector(code_name, string):
    detector = re.compile(CODES_REGEX[code_name], re.I)
    detected = detector.findall(string)
    detected_list = list(sum(detected, ()))
    clean_list = []
    for element in detected_list:
        if element != "" and reformat_results(element) not in clean_list:
            clean_list.append(reformat_results(element))
    sorted_list = sorted(clean_list)
    dict_list = []
    for element in sorted_list:
        dict_entry = {"number": element}
        dict_list.append(dict_entry)
    return dict_list


# Les numéros d'articles du texte sont reformatés pour en retirer certains
# caractères (espace, espace insécable, point) qui empêchent la recherche Légifrance.
def reformat_results(result):
    new_result = ""
    result = result.replace("\xa0", " ")
    if (result[0]).isdigit():
        return result
    else:
        for char in result:
            if char != "." and char != " ":
                new_result = new_result + char
    return new_result


# Recherche sur Légifrance de l'identifiant unique de l'article
def get_article_id(article_number, code_name):
    data = {
        "recherche": {
            "champs": [
                {
                    "typeChamp": "NUM_ARTICLE",
                    "criteres": [
                        {
                            "typeRecherche": "EXACTE",
                            "valeur": article_number,
                            "operateur": "ET",
                        }
                    ],
                    "operateur": "ET",
                }
            ],
            "filtres": [
                {"facette": "NOM_CODE", "valeurs": [code_name]},
                {"facette": "DATE_VERSION", "singleDate": today},
            ],
            "pageNumber": 1,
            "pageSize": 10,
            "operateur": "ET",
            "sort": "PERTINENCE",
            "typePagination": "ARTICLE",
        },
        "fond": "CODE_DATE",
    }

    response = session.post(
        "https://api.piste.gouv.fr/dila/legifrance-beta/lf-engine-app/search",
        headers=headers,
        json=data,
    )

    article_informations = json.loads(response.text)
    if not article_informations["results"]:
        return None
    else:
        article_id = article_informations["results"][0]["sections"][0]["extracts"][0][
            "id"
        ]
        return article_id


# À partir de l'identifiant unique, rapatriement du contenu de l'article
def get_article_content(article_id):
    data = {"id": article_id}
    response = session.post(
        "https://api.piste.gouv.fr/dila/legifrance-beta/lf-engine-app/consult/getArticle",
        headers=headers,
        json=data,
    )
    article_dictionary = json.loads(response.text)
    return article_dictionary["article"]


# Légifrance utilise des dates au format Epoch qu'il faut convertir au format classique
def epoch_converter(epoch):
    converted_date = datetime.datetime(1970, 1, 1) + datetime.timedelta(
        seconds=(epoch / 1000)
    )
    simplified_date = converted_date.strftime("%d-%m-%Y")
    return simplified_date


# Initialisation du programme



def validity_period(user_past=3, user_future=3):
    """Définition des bornes de la période déclenchant une alerte
    si l'article a été modifié / va être modifié
    """
    past_reference = (
                             datetime.datetime.now() - datetime.timedelta(days=float(user_past) * 365)
                     ).timestamp() * 1000
    future_reference = (
                               datetime.datetime.now() + datetime.timedelta(days=float(user_future) * 365)
                       ).timestamp() * 1000
    return (past_reference, future_reference)

def upload_document(upload="./tests/docs/newtest.docx"):
    # upload = request.files.get("upload")
    if upload is None:
        raise Exception("Pas de fichier.")
        
    doc_name, doc_ext = os.path.splitext(upload.filename)
    if doc_ext not in (".odt", ".docx", ".pdf"):
        raise Exception("Extension incorrecte: les fichiers acceptés terminent par *.odt, *docx, *.pdf")
        
    save_path = Path.cwd() / Path("tmp")
    if not os.path.exists(save_path):
        os.makedirs(save_path)
    file_path = "{path}/{file}".format(path=save_path, file=upload.filename)
    upload.save(file_path, overwrite="true")
    file_size = os.stat(file_path).st_size
    if file_size > 2000000:
        raise Exception("Fichier trop lourd: la taille du fichier doit être < à 2Mo")
    return (doc_name, doc_ext, file_path)


# Affichage de la page web d'accueil

@app.route("/")
def root():
    return static_file("index.html", root=".")



# Actions à effectuer à l'upload du document de l'utilisateur
@app.route("/upload", method="POST")
def do_upload():
    # global headers

    # load_dotenv(find_dotenv())
    # headers = legifrance_auth()

    # password = os.environ.get("PASSWORD")
    # user_password = request.forms.get("password")
    # if user_password != password:
    #     yield "Mot de passe incorrect"
    #     sys.exit()
    
    # CODE RESULTS
    code_results = {}
    for code in MAIN_CODELIST:
        code_results[code] = []

    # L'utilisateur définit sur quelle période la validité de l'article est testée
    # user_past = request.forms.get("user_past")
    # user_future = request.forms.get("user_future")
    user_past,user_future = validity_period(request.forms.get("user_past"), request.forms.get("user_future"))
    

    # L'utilisateur upload son document, il est enregistré provisoirement
    doc_name, doc_ext, file_path = upload_document(request.files.get("upload"))
    yield "<!DOCTYPE html>"
    yield "<head>"
    yield """<title> Code is low</title>"""
    yield """<link rel="stylesheet" href="https://www.w3schools.com/w3css/4/w3.css">"""
    yield """</head>"""
    yield """<body>"""
    yield """<div class="w3-container w3-blue-grey">"""
    yield """<h1> Code is low</h1>"""
    yield """</div>"""
    yield "<h3> Analyse en cours. Veuillez patienter... </h3>"
    yield "<h4> Le fichier " + doc_name + doc_ext.upper() + " est actuellement parcouru. </h4>"

    cleantext = yield from file_opener(doc_ext, file_path)

    # Suppression du fichier utilisateur, devenu inutile
    os.remove(file_path)

    # Mise en œuvre des expressions régulières
    yield "<h4> Les différents codes de droit français sont recherchés. </h4>"
    for code_name in MAIN_CODELIST:
        yield "<small> " + code_name + "...  </small>"
        code_results[code_name] = code_detector(code_name, cleantext)

    for code_name in MAIN_CODELIST:
        if not code_results[code_name]:
            del code_results[code_name]

    # Récupération des caractéristiques de chaque article sur Légifrance
    for code_name in code_results:
        yield "<h4> " + "La base Légifrance est interrogée : textes du " + MAIN_CODELIST[code_name] + "... </h4>"
        for article in code_results[code_name]:
            yield "<small> " + "Article " + article["number"] + "...  </small>"
            article_id = get_article_id(article["number"], MAIN_CODELIST[code_name])
            article.update({"id": article_id})

            if article["id"] is not None:
                article_content = get_article_content(article["id"])
                # article_text = article_content["texte"]
                article_start = article_content["dateDebut"]
                article_end = article_content["dateFin"]
                article.update({"start": article_start, "end": article_end})

    # Tri des articles pour affichage final
    yield "<h5> Tri des résultats en cours. </h5>"
    articles_not_found = []
    articles_recently_modified = []
    articles_changing_soon = []
    articles_without_event = []

    for code_name in code_results:
        for article in code_results[code_name]:
            if article["id"] is None:
                articles_not_found.append(
                    "Article " + article["number"] + " du " + MAIN_CODELIST[code_name] + " non trouvé"
                )
            else:
                article_hyperlink = """<a class="w3-text-blue" href="https://www.legifrance.gouv.fr/codes/article_lc/""" + \
                                    article[
                                        "id"] + """" target="_blank" rel="noopener">""" + article["number"] + """</a>"""

                if article["start"] < past_reference and article["end"] > future_reference:
                    articles_without_event.append(
                        "Article " + article_hyperlink + " du " + MAIN_CODELIST[code_name]
                    )
                if article["start"] > past_reference:
                    articles_recently_modified.append(
                        "L'article "
                        + article_hyperlink
                        + " du "
                        + MAIN_CODELIST[code_name]
                        + " a été modifié le "
                        + str(epoch_converter(article["start"]))
                    )
                if article["end"] < future_reference:
                    articles_changing_soon.append(
                        "La version actuelle de l'article "
                        + article_hyperlink
                        + " du "
                        + MAIN_CODELIST[code_name]
                        + " n'est valable que jusqu'au "
                        + str(epoch_converter(article["end"]))
                    )

    # Utilisation d'un template Bottle pour afficher les résultats
    yield template(
        "results",
        **{
            "articles_not_found": articles_not_found,
            "articles_recently_modified": articles_recently_modified,
            "articles_changing_soon": articles_changing_soon,
            "articles_without_event": articles_without_event,
            "user_past": user_past,
            "user_future": user_future,
        },
    )


# Corps du programme

if __name__ == "__main__":
    today = int(time.time() * 1000)

    retry_strategy = Retry(
        total=3,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["HEAD", "GET", "POST", "OPTIONS"],
    )
    adapter = HTTPAdapter(max_retries=retry_strategy)
    session = requests.Session()
    session.mount("https://", adapter)
    session.mount("http://", adapter)

if os.environ.get("APP_LOCATION") == "heroku":
    SSLify(app)
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
else:
    app.run(host="localhost", port=8080, debug=True, reloader=True)
